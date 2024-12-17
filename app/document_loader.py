import os

from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


# Retrieve OpenAI API key from system environment variable
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

index = None

def load_documents_from_data():
     """
     Load documents from the 'data' directory, supporting .txt, .pdf, and .docx files.
     """
     documents = []
     data_dir = 'data'

     for filename in os.listdir(data_dir):
          filepath = os.path.join(data_dir,filename)

          if filename.endswith(".txt"):
               with open(filepath, 'r', encoding='utf-8') as f:
                    documents.append(f.read())
          elif filename.endswith(".pdf"):
               reader = PdfReader(filepath)
               text = " ".join([page.extract_text() for page in reader.pages])
               documents.append(text)
          elif filename.endswith(".docx"):
               doc = DocxDocument(filepath)
               text = " ".join([paragraph.text for paragraph in doc.paragraphs])
               documents.append(text)

     return documents


def build_index():
     """
     Load documents from the 'data' directory and create a vector store index.
     """
     global index

     # Load documents
     raw_documents = load_documents_from_data()

     # Convert raw text into LlamaIndex Documents
     documents = [Document(text=doc) for doc in raw_documents]

     #Create a vector store index
     index = VectorStoreIndex.from_documents(documents)

     return index

build_index()